import os
import re
from pathlib import Path
from uuid import uuid4

from dotenv import load_dotenv
from pypdf import PdfReader
from docx import Document

from qdrant_client.models import Distance, VectorParams, PointStruct

# Reutilizamos la conexión y los embeddings centralizados en client.py
# para no duplicar la lógica de OpenAI/Qdrant.
from app.rag.client import (
    OPENAI_EMBEDDING_MODEL,
    crear_embedding,
    obtener_cliente_qdrant,
)


load_dotenv()

QDRANT_URL = os.getenv("QDRANT_URL")
QDRANT_API_KEY = os.getenv("QDRANT_API_KEY")
COLLECTION = os.getenv("QDRANT_COLLECTION", "palermia_docs")

DOCS_DIR = Path("docs_rag")


def extraer_texto_pdf(path_pdf: Path) -> str:
    reader = PdfReader(str(path_pdf))
    texto = ""

    for i, page in enumerate(reader.pages):
        page_text = page.extract_text()
        if page_text:
            texto += f"\n\n--- Página {i + 1} ---\n{page_text}"

    return texto


def extraer_texto_docx(path_docx: Path) -> str:
    document = Document(str(path_docx))
    partes = []

    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            partes.append(paragraph.text.strip())

    for table in document.tables:
        for row in table.rows:
            celdas = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if celdas:
                partes.append(" | ".join(celdas))

    return "\n".join(partes)


def extraer_texto_documento(path: Path) -> str:
    if path.suffix.lower() == ".pdf":
        return extraer_texto_pdf(path)

    if path.suffix.lower() == ".docx":
        return extraer_texto_docx(path)

    raise ValueError(f"Formato no soportado: {path.name}")

def dividir_por_articulos(texto: str) -> list[str]:
    """
    Divide un texto normativo tomando como separador la palabra 'Artículo'.

    Ejemplos detectados:
    - Artículo 1.
    - Artículo 10. Autoridad competente
    - ARTÍCULO 4
    - Articulo 7

    La idea es que cada artículo quede como una unidad semántica propia.
    """

    patron = r"(?i)(?=art[íi]culo\s+\d+)"

    partes = re.split(patron, texto)

    articulos = []

    for parte in partes:
        parte = parte.strip()

        if not parte:
            continue

        # Nos quedamos sólo con bloques que parecen artículos.
        if re.match(r"(?i)^art[íi]culo\s+\d+", parte):
            articulos.append(parte)

    return articulos


def subdividir_texto_largo(
    texto: str,
    chunk_size: int = 900,
    overlap: int = 150,
) -> list[str]:
    """
    Subdivide un texto largo en chunks por caracteres.

    Esto se usa sólo cuando:
    - el documento no tiene artículos;
    - o un artículo es demasiado largo.
    """

    chunks = []
    inicio = 0

    while inicio < len(texto):
        fin = inicio + chunk_size
        chunk = texto[inicio:fin].strip()

        if chunk:
            chunks.append(chunk)

        inicio += chunk_size - overlap

    return chunks



def crear_chunks(
    texto: str,
    max_articulo_size: int = 1200,
    fallback_chunk_size: int = 700,
    fallback_overlap: int = 120,
) -> list[str]:
    """
    Crea chunks inteligentes para documentos institucionales.

    Estrategia:

    1. Primero intenta dividir por artículos.
       Esto es ideal para manuales, reglamentos y normas internas.

    2. Si encuentra artículos:
       - cada artículo corto queda como un chunk;
       - cada artículo largo se subdivide para no generar embeddings enormes.

    3. Si NO encuentra artículos:
       - usa chunking tradicional por caracteres.
       Esto sirve para actas, resoluciones, políticas o documentos narrativos.
    """

    articulos = dividir_por_articulos(texto)

    chunks = []

    if articulos:
        print(f"Se detectaron {len(articulos)} artículos. Chunking por artículo.")

        for articulo in articulos:
            if len(articulo) <= max_articulo_size:
                chunks.append(articulo)
            else:
                subchunks = subdividir_texto_largo(
                    articulo,
                    chunk_size=fallback_chunk_size,
                    overlap=fallback_overlap,
                )
                chunks.extend(subchunks)

        return chunks

    print("No se detectaron artículos. Chunking por caracteres.")

    return subdividir_texto_largo(
        texto,
        chunk_size=fallback_chunk_size,
        overlap=fallback_overlap,
    )


def listar_documentos():
    return sorted(
        [
            path
            for path in DOCS_DIR.iterdir()
            if path.suffix.lower() in [".pdf", ".docx"]
        ]
    )


def main():
    print("=" * 80)
    print("INGESTA RAG - PALERMIA S.A. / OPENAI EMBEDDINGS")
    print("=" * 80)

    if not QDRANT_URL:
        raise ValueError("Falta QDRANT_URL en .env")

    if not QDRANT_API_KEY:
        raise ValueError("Falta QDRANT_API_KEY en .env")

    if not os.getenv("OPENAI_API_KEY"):
        raise ValueError("Falta OPENAI_API_KEY en .env")

    if not DOCS_DIR.exists():
        raise FileNotFoundError("No existe la carpeta docs_rag")

    documentos = listar_documentos()

    if not documentos:
        print("No hay documentos PDF o DOCX en docs_rag/")
        return

    qdrant_client = obtener_cliente_qdrant()

    embedding_prueba = crear_embedding("texto de prueba")
    vector_size = len(embedding_prueba)

    print(f"Modelo embeddings: {OPENAI_EMBEDDING_MODEL}")
    print(f"Dimensión vectorial: {vector_size}")
    print(f"Collection: {COLLECTION}")

    if not qdrant_client.collection_exists(collection_name=COLLECTION):
        print("Collection no existe. Creando...")

        qdrant_client.create_collection(
            collection_name=COLLECTION,
            vectors_config=VectorParams(
                size=vector_size,
                distance=Distance.COSINE,
            ),
        )

        print("Collection creada.")
    else:
        print("Collection existente. Se agregan documentos.")

    points = []

    for documento in documentos:
        print(f"\nProcesando: {documento.name}")

        texto = extraer_texto_documento(documento)

        if not texto.strip():
            print("No se pudo extraer texto. Se omite.")
            continue

        chunks = crear_chunks(texto)
        print(f"Chunks generados: {len(chunks)}")

        for index, chunk in enumerate(chunks):
            vector = crear_embedding(chunk)

            points.append(
                PointStruct(
                    id=str(uuid4()),
                    vector=vector,
                    payload={
                        "text": chunk,
                        "source": documento.name,
                        "chunk": index,
                        "extension": documento.suffix.lower(),
                    },
                )
            )

    if not points:
        print("No se generaron puntos para subir.")
        return

    print("\nSubiendo vectores a Qdrant...")

    qdrant_client.upsert(
        collection_name=COLLECTION,
        points=points,
    )

    print("\nIngesta finalizada.")
    print(f"Documentos procesados: {len(documentos)}")
    print(f"Chunks subidos: {len(points)}")


if __name__ == "__main__":
    main()